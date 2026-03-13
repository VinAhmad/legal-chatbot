"""Service untuk export jawaban chatbot ke DOCX dan PDF."""

import io
import re
import logging
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from fpdf import FPDF

logger = logging.getLogger(__name__)


def _parse_markdown_blocks(text: str) -> list[dict]:
    """
    Parse markdown sederhana jadi list blok terstruktur.
    Support: heading, bold, list, table, paragraph.
    """
    lines = text.split("\n")
    blocks: list[dict] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,3})\s+(.+)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2)})
            i += 1
            continue

        # Table: kumpulkan semua baris table
        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "table", "lines": table_lines})
            continue

        # Unordered list
        list_match = re.match(r"^[-*]\s+(.+)", stripped)
        if list_match:
            items = []
            while i < len(lines):
                lm = re.match(r"^\s*[-*]\s+(.+)", lines[i])
                if not lm:
                    break
                items.append(lm.group(1))
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # Ordered list
        ol_match = re.match(r"^\d+[.)]\s+(.+)", stripped)
        if ol_match:
            items = []
            while i < len(lines):
                om = re.match(r"^\s*\d+[.)]\s+(.+)", lines[i])
                if not om:
                    break
                items.append(om.group(1))
                i += 1
            blocks.append({"type": "ordered_list", "items": items})
            continue

        # Paragraph biasa (gabungkan baris yang berdekatan)
        para_lines = []
        while i < len(lines):
            cl = lines[i].strip()
            if not cl or cl.startswith("#") or cl.startswith("|") or re.match(r"^[-*]\s+", cl) or re.match(r"^\d+[.)]\s+", cl):
                break
            para_lines.append(cl)
            i += 1
        if para_lines:
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def _parse_table_lines(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse markdown table lines jadi headers dan rows."""
    if len(lines) < 2:
        return [], []

    def split_cells(line: str) -> list[str]:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        return cells

    headers = split_cells(lines[0])
    rows = []
    for line in lines[1:]:
        # Skip separator line (|---|---|)
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        rows.append(split_cells(line))
    return headers, rows


def _add_rich_text(paragraph, text: str):
    """Tambahkan teks dengan parsing bold (**text**) ke paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def export_to_docx(content: str, sources: list[dict] | None = None) -> bytes:
    """Convert markdown content ke DOCX bytes."""
    doc = DocxDocument()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Hasil Analisis Dokumen Legal")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(30, 58, 138)

    # Tanggal
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Diekspor: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph("")

    blocks = _parse_markdown_blocks(content)

    for block in blocks:
        if block["type"] == "heading":
            level = min(block["level"], 3)
            doc.add_heading(block["text"], level=level)

        elif block["type"] == "paragraph":
            para = doc.add_paragraph()
            _add_rich_text(para, block["text"])

        elif block["type"] == "list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Bullet")
                _add_rich_text(para, item)

        elif block["type"] == "ordered_list":
            for item in block["items"]:
                para = doc.add_paragraph(style="List Number")
                _add_rich_text(para, item)

        elif block["type"] == "table":
            headers, rows = _parse_table_lines(block["lines"])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ci, h in enumerate(headers):
                    cell = table.rows[0].cells[ci]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)

                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < len(headers):
                            cell = table.rows[ri + 1].cells[ci]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                doc.add_paragraph("")

    # Sumber referensi
    if sources:
        doc.add_paragraph("")
        doc.add_heading("Sumber Referensi", level=2)
        for idx, src in enumerate(sources, 1):
            name = src.get("document_name", "Unknown")
            page = src.get("page")
            page_str = f", Hal. {page}" if page else ""
            doc.add_paragraph(f"{idx}. {name}{page_str}", style="List Number")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


class _LegalPDF(FPDF):
    """PDF custom dengan header dan footer."""

    def header(self):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(107, 114, 128)
        self.cell(0, 8, "Hasil Analisis - Legal AI Assistant", align="L")
        self.ln(10)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(156, 163, 175)
        self.cell(0, 10, f"Halaman {self.page_no()}/{{nb}}", align="C")


def export_to_pdf(content: str, sources: list[dict] | None = None) -> bytes:
    """Convert markdown content ke PDF bytes."""
    pdf = _LegalPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(30, 58, 138)
    pdf.cell(0, 12, "Hasil Analisis Dokumen Legal", align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(0, 8, f"Diekspor: {datetime.now().strftime('%d %B %Y, %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    blocks = _parse_markdown_blocks(content)

    for block in blocks:
        if block["type"] == "heading":
            size = {1: 16, 2: 13, 3: 11}.get(block["level"], 11)
            pdf.set_font("Helvetica", "B", size)
            pdf.set_text_color(30, 41, 59)
            pdf.ln(3)
            pdf.multi_cell(0, 7, block["text"])
            pdf.ln(2)

        elif block["type"] == "paragraph":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(55, 65, 81)
            clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", block["text"])
            pdf.multi_cell(0, 6, clean)
            pdf.ln(3)

        elif block["type"] == "list":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(55, 65, 81)
            for item in block["items"]:
                clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
                pdf.cell(6)
                pdf.cell(4, 6, chr(8226))
                pdf.multi_cell(0, 6, f" {clean}")
            pdf.ln(2)

        elif block["type"] == "ordered_list":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(55, 65, 81)
            for idx, item in enumerate(block["items"], 1):
                clean = re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
                pdf.cell(6)
                pdf.cell(8, 6, f"{idx}.")
                pdf.multi_cell(0, 6, f" {clean}")
            pdf.ln(2)

        elif block["type"] == "table":
            headers, rows = _parse_table_lines(block["lines"])
            if headers:
                col_count = len(headers)
                available_width = pdf.w - 20
                col_width = available_width / col_count

                # Header row
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_fill_color(239, 246, 255)
                pdf.set_text_color(30, 58, 138)
                for h in headers:
                    pdf.cell(col_width, 8, h[:30], border=1, fill=True, align="C")
                pdf.ln()

                # Data rows
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(55, 65, 81)
                for ri, row_data in enumerate(rows):
                    if ri % 2 == 1:
                        pdf.set_fill_color(249, 250, 251)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    for ci in range(col_count):
                        cell_text = row_data[ci] if ci < len(row_data) else ""
                        pdf.cell(col_width, 7, cell_text[:40], border=1, fill=True)
                    pdf.ln()
                pdf.ln(4)

    # Sumber referensi
    if sources:
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(30, 41, 59)
        pdf.cell(0, 8, "Sumber Referensi", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(55, 65, 81)
        for idx, src in enumerate(sources, 1):
            name = src.get("document_name", "Unknown")
            page = src.get("page")
            page_str = f", Hal. {page}" if page else ""
            pdf.cell(0, 6, f"{idx}. {name}{page_str}", new_x="LMARGIN", new_y="NEXT")

    return pdf.output()
