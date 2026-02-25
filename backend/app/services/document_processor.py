"""Parse berbagai format dokumen (PDF, DOCX, CSV, TXT) menjadi text chunks."""

import os
from pathlib import Path
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd
from langchain_text_splitters import RecursiveCharacterTextSplitter


@dataclass
class DocumentChunk:
    text: str
    metadata: dict


CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
    separators=["\n\n", "\n", ". ", " ", ""],
)


class DocumentProcessor:
    """Parse berbagai format dokumen menjadi list of text chunks."""

    @staticmethod
    async def process(file_path: str, filename: str) -> list[DocumentChunk]:
        ext = Path(file_path).suffix.lower()
        parsers = {
            ".pdf": DocumentProcessor._parse_pdf,
            ".docx": DocumentProcessor._parse_docx,
            ".csv": DocumentProcessor._parse_csv,
            ".xlsx": DocumentProcessor._parse_excel,
            ".xls": DocumentProcessor._parse_excel,
            ".txt": DocumentProcessor._parse_text,
            ".md": DocumentProcessor._parse_text,
        }

        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"Format file tidak didukung: {ext}")

        raw_texts = parser(file_path)
        chunks = DocumentProcessor._split_into_chunks(raw_texts, filename)
        return chunks

    @staticmethod
    def _parse_pdf(file_path: str) -> list[dict]:
        """Extract teks dari setiap halaman PDF."""
        results = []
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text").strip()
            if text:
                results.append({"text": text, "page": page_num + 1})
        doc.close()
        return results

    @staticmethod
    def _parse_docx(file_path: str) -> list[dict]:
        """Extract teks dari paragraf dan tabel di DOCX."""
        doc = DocxDocument(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        combined = "\n".join(full_text)
        return [{"text": combined, "page": None}] if combined else []

    @staticmethod
    def _parse_csv(file_path: str) -> list[dict]:
        """Parse CSV jadi teks per baris."""
        df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="skip")
        return DocumentProcessor._dataframe_to_texts(df)

    @staticmethod
    def _parse_excel(file_path: str) -> list[dict]:
        """Parse Excel jadi teks per sheet."""
        results = []
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            sheet_texts = DocumentProcessor._dataframe_to_texts(df, sheet_name)
            results.extend(sheet_texts)
        return results

    @staticmethod
    def _dataframe_to_texts(df: pd.DataFrame, sheet_name: str | None = None) -> list[dict]:
        """Convert DataFrame rows jadi list of text dicts."""
        results = []
        headers = list(df.columns)
        header_line = " | ".join(str(h) for h in headers)

        for idx, row in df.iterrows():
            row_text = " | ".join(f"{h}: {row[h]}" for h in headers if pd.notna(row[h]))
            if row_text:
                full_text = f"[Header: {header_line}]\n{row_text}"
                meta = {"page": None}
                if sheet_name:
                    meta["sheet"] = sheet_name
                results.append({"text": full_text, **meta})

        return results

    @staticmethod
    def _parse_text(file_path: str) -> list[dict]:
        """Parse plain text / markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        return [{"text": content, "page": None}] if content else []

    @staticmethod
    def _split_into_chunks(raw_texts: list[dict], filename: str) -> list[DocumentChunk]:
        """Split raw texts jadi chunks yang lebih kecil pakai RecursiveCharacterTextSplitter."""
        chunks = []
        for item in raw_texts:
            text = item["text"]
            page = item.get("page")
            sheet = item.get("sheet")

            splits = text_splitter.split_text(text)
            for split_text in splits:
                metadata = {"source": filename}
                if page is not None:
                    metadata["page"] = page
                if sheet is not None:
                    metadata["sheet"] = sheet
                chunks.append(DocumentChunk(text=split_text, metadata=metadata))

        return chunks
